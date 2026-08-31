import io
import re
from difflib import SequenceMatcher
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

class WordProcessor:
    @staticmethod
    def _clean_str(text: str) -> str:
        """Chuẩn hóa chuỗi: bỏ khoảng trắng thừa, ký tự đặc biệt để so khớp chính xác hơn."""
        if not text:
            return ""
        # Thay thế ký tự khoảng trắng đặc biệt thành dấu cách chuẩn
        text = text.replace('\xa0', ' ').replace('\t', ' ').replace('\r', '')
        # Bỏ các ký tự markdown nếu AI vô tình thêm vào
        text = re.sub(r'[*_#`]', '', text)
        return re.sub(r'\s+', ' ', text).strip().lower()

    @staticmethod
    def extract_text(file_bytes: bytes) -> str:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            full_text.append(p.text.strip())
        return "\n".join(full_text)

    @staticmethod
    def insert_paragraph_after(paragraph, text, color_rgb, prefix=""):
        """Chèn đoạn văn mới liền sau paragraph chỉ định."""
        new_p = OxmlElement('w:p')
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        new_para.paragraph_format.space_before = Pt(3)
        new_para.paragraph_format.space_after = Pt(4)
        new_para.paragraph_format.line_spacing = 1.15
        
        # Thêm tiền tố in đậm nhẹ
        if prefix:
            run_p = new_para.add_run(f"{prefix} ")
            run_p.font.color.rgb = color_rgb
            run_p.bold = True
            run_p.italic = True
            run_p.font.size = Pt(11)

        run_t = new_para.add_run(text)
        run_t.font.color.rgb = color_rgb
        run_t.italic = True
        run_t.font.size = Pt(11)
        return new_para

    @staticmethod
    def integrate_digital_capacity(file_bytes: bytes, ai_data: dict, integration_type: str) -> io.BytesIO:
        doc = Document(io.BytesIO(file_bytes))
        sua_doi_list = ai_data.get('sua_doi', [])
        
        color_digital = RGBColor(0, 102, 204)   # Xanh dương
        color_ai = RGBColor(214, 107, 0)        # Vàng cam
        
        # Thu thập toàn bộ paragraphs từ cả văn bản chính và các bảng
        all_paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        all_paragraphs.append(p)

        used_paragraphs = set()

        for item in sua_doi_list:
            raw_anchor = item.get('anchor_text', '').strip()
            content = item.get('insert_content', '').strip()
            loai = item.get('loai', 'Năng lực số')
            
            if not raw_anchor or not content:
                continue
            
            clean_anchor = WordProcessor._clean_str(raw_anchor)
            prefix = "[Năng lực AI]:" if loai == "Năng lực AI" else "[Năng lực số]:"
            color = color_ai if loai == "Năng lực AI" else color_digital
            
            inserted = False
            best_match_para = None
            best_ratio = 0.0

            # BƯỚC 1: Tìm kiếm chính xác hoặc chứa chuỗi (Sub-string match)
            for para in all_paragraphs:
                clean_p_text = WordProcessor._clean_str(para.text)
                if not clean_p_text:
                    continue
                
                # Kiểm tra chuỗi chứa nhau
                if (clean_anchor in clean_p_text or clean_p_text in clean_anchor) and para not in used_paragraphs:
                    WordProcessor.insert_paragraph_after(para, content, color, prefix)
                    used_paragraphs.add(para)
                    inserted = True
                    break
                
                # Tính toán độ tương đồng phòng khi không khớp 100%
                ratio = SequenceMatcher(None, clean_anchor, clean_p_text).ratio()
                if ratio > best_ratio:
                    best_ratio = ratio
                    best_match_para = para

            # BƯỚC 2: Fallback - Nếu không khớp tuyệt đối, chèn vào đoạn có độ tương đồng cao nhất (>= 55%)
            if not inserted and best_match_para is not None and best_ratio >= 0.55:
                if best_match_para not in used_paragraphs:
                    WordProcessor.insert_paragraph_after(best_match_para, content, color, prefix)
                    used_paragraphs.add(best_match_para)
                    inserted = True

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        return output_stream
